from pathlib import Path
import base64
import hashlib
import io
import json
import math
import mimetypes
import re
import shutil
import zipfile

from sqlalchemy import func, or_
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.entities import AppSetting, Document, DocumentChunk, DocumentChunkEmbedding, DocumentRevision, DocumentTask


def document_workspace(document_id: int) -> Path:
    return get_settings().knowledge_dir / str(document_id)


def document_assets_dir(document_id: int) -> Path:
    return document_workspace(document_id) / "assets"


def document_markdown_path(document_id: int) -> Path:
    return document_workspace(document_id) / "content.md"


def read_document_markdown(document_id: int) -> str:
    path = document_markdown_path(document_id)
    if path.exists():
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def markdown_for_api(document_id: int, markdown: str) -> str:
    """把本地 assets 链接改成 data URL，避免前端图片请求丢失 Authorization."""

    def replace(match: re.Match) -> str:
        alt, asset_name = match.group(1), Path(match.group(2)).name
        asset_path = document_assets_dir(document_id) / asset_name
        if not asset_path.exists():
            return match.group(0)
        mime = mimetypes.guess_type(asset_path.name)[0] or "application/octet-stream"
        data = base64.b64encode(asset_path.read_bytes()).decode("ascii")
        return f"![{alt}](data:{mime};base64,{data})"

    return re.sub(r"!\[([^\]]*)\]\(assets/([^)]+)\)", replace, markdown)


def delete_document_files(document_id: int) -> None:
    shutil.rmtree(document_workspace(document_id), ignore_errors=True)


def build_markdown_document(path: Path, document_id: int, db: Session | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix == ".md":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        text = extract_text(path, db)
        return clean_text(text)
    if suffix == ".docx":
        return extract_docx_markdown(path, document_id, db)
    raise RuntimeError(f"不支持的文档格式: {suffix}")


def extract_text(path: Path, db: Session | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pypdf 未安装，无法解析 PDF") from exc
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        return extract_docx_text(path, db)
    raise RuntimeError(f"不支持的文档格式: {suffix}")


def extract_docx_markdown(path: Path, document_id: int, db: Session | None = None) -> str:
    try:
        from docx import Document as DocxDocument
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
    except ImportError as exc:
        raise RuntimeError("python-docx 未安装，无法解析 DOCX") from exc

    doc = DocxDocument(str(path))
    assets_dir = document_assets_dir(document_id)
    shutil.rmtree(assets_dir, ignore_errors=True)
    assets_dir.mkdir(parents=True, exist_ok=True)

    image_state = {"index": 0, "vision_used": 0}
    blocks: list[str] = []
    parent = doc
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            markdown = _paragraph_to_markdown(Paragraph(child, parent), doc, assets_dir, image_state, db)
        elif isinstance(child, CT_Tbl):
            markdown = _table_to_markdown(Table(child, parent))
        else:
            markdown = ""
        if markdown:
            blocks.append(markdown)
    return "\n\n".join(blocks).strip()


def _paragraph_to_markdown(paragraph, doc, assets_dir: Path, image_state: dict, db: Session | None) -> str:
    inline_parts: list[str] = []
    for run in paragraph.runs:
        if run.text:
            inline_parts.append(run.text)
        inline_parts.extend(_run_images_to_markdown(run, doc, assets_dir, image_state, db))

    text = clean_text("".join(inline_parts))
    if not text:
        return ""

    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if style_name.startswith("heading"):
        level_match = re.search(r"(\d+)", style_name)
        level = min(max(int(level_match.group(1)) if level_match else 2, 1), 4)
        return f"{'#' * level} {text.lstrip('#').strip()}"
    if "list bullet" in style_name:
        return f"- {text}"
    if "list number" in style_name:
        return f"1. {text}"
    return text


def _run_images_to_markdown(run, doc, assets_dir: Path, image_state: dict, db: Session | None) -> list[str]:
    image_markdown = []
    blips = run._element.xpath(".//*[local-name()='blip']")
    for blip in blips:
        embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not embed or embed not in doc.part.related_parts:
            continue
        part = doc.part.related_parts[embed]
        image_state["index"] += 1
        image_index = image_state["index"]
        ext = Path(str(part.partname)).suffix.lower() or _content_type_ext(getattr(part, "content_type", ""))
        if ext not in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff"}:
            ext = ".png"
        asset_name = f"image-{image_index:03d}{ext}"
        image_bytes = part.blob
        (assets_dir / asset_name).write_bytes(image_bytes)
        image_markdown.append(f"\n\n![图片 {image_index}](assets/{asset_name})")
        image_text = _image_text_for_markdown(asset_name, image_bytes, image_state, db)
        if image_text:
            image_markdown.append(f"\n\n<details><summary>图片 {image_index} 识别文本</summary>\n\n{image_text}\n\n</details>")
    return image_markdown


def _content_type_ext(content_type: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/webp": ".webp",
        "image/tiff": ".tiff",
    }
    return mapping.get(content_type, ".png")


def _image_text_for_markdown(image_name: str, image_bytes: bytes, image_state: dict, db: Session | None) -> str:
    api_key, base_url, model = _resolve_vision_config(db)
    env = get_settings()
    if api_key and image_state["vision_used"] < max(0, env.docx_vision_max_images):
        image_state["vision_used"] += 1
        return describe_image_with_vision_model(image_name, image_bytes, api_key, base_url, model)
    if not api_key:
        return describe_image_with_ocr(image_bytes)
    return ""


def image_text_for_markdown(image_name: str, image_bytes: bytes, image_state: dict, db: Session | None) -> str:
    return _image_text_for_markdown(image_name, image_bytes, image_state, db)


def describe_image_with_vision_model(image_name: str, image_bytes: bytes, api_key: str, base_url: str | None, model: str) -> str:
    try:
        from openai import OpenAI
    except ImportError:
        return describe_image_with_ocr(image_bytes)
    try:
        client = OpenAI(api_key=api_key, base_url=base_url)
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": "你是运维文档解析助手。请读取图片中的文字、命令、错误信息、配置、表格和架构关系，不要编造。",
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "请解析这张运维文档图片，输出简洁 Markdown。看不清时直接说明。"},
                        {"type": "image_url", "image_url": {"url": image_to_data_url(image_name, image_bytes)}},
                    ],
                },
            ],
        )
        return (completion.choices[0].message.content or "").strip()
    except Exception:
        return ""


def describe_image_with_ocr(image_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""
    if not tesseract_available():
        return ""
    try:
        image = Image.open(io.BytesIO(image_bytes))
        return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
    except Exception:
        return ""


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [clean_text(cell.text).replace("|", "\\|") for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def extract_docx_text(path: Path, db: Session | None = None) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx 未安装，无法解析 DOCX") from exc

    doc = DocxDocument(str(path))
    parts: list[str] = []

    paragraph_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraph_text:
        parts.append("## 正文\n" + "\n".join(paragraph_text))

    table_blocks = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            table_blocks.append(f"### 表格 {table_index}\n" + "\n".join(rows))
    if table_blocks:
        parts.append("## 表格\n" + "\n\n".join(table_blocks))

    image_text, image_note = extract_docx_images_text(path, db)
    if image_text:
        parts.append("## 图片解析\n" + image_text)
    if image_note:
        parts.append("## 图片解析提示\n" + image_note)

    return "\n\n".join(parts)


def extract_docx_images_text(path: Path, db: Session | None = None) -> tuple[str, str]:
    image_names = list_docx_images(path)
    if not image_names:
        return "", ""

    api_key, base_url, model = _resolve_vision_config(db)
    if api_key:
        return extract_docx_images_with_vision_model(path, image_names, api_key, base_url, model)

    return extract_docx_images_with_ocr(path, image_names)


def _resolve_vision_config(db: Session | None) -> tuple[str, str, str]:
    """解析视觉模型配置：DB 优先，env 兜底."""
    env = get_settings()
    api_key = env.openai_api_key or ""
    base_url = env.openai_base_url or None
    model = env.openai_vision_model or env.openai_model

    if db:
        from app.models.entities import AppSetting  # noqa: F811
        keys = ["optimize_api_key", "optimize_base_url", "optimize_vision_model"]
        db_vals = {}
        for row in db.query(AppSetting).filter(AppSetting.key.in_(keys)).all():
            db_vals[row.key] = row.value
        api_key = db_vals.get("optimize_api_key") or api_key
        base_url = db_vals.get("optimize_base_url") or base_url
        model = db_vals.get("optimize_vision_model") or model

    return api_key, base_url, model


def extract_docx_images_with_vision_model(path: Path, image_names: list[str], api_key: str | None = None, base_url: str | None = None, model: str | None = None) -> tuple[str, str]:
    env = get_settings()
    api_key = api_key or env.openai_api_key
    base_url = base_url or env.openai_base_url
    model = model or env.openai_vision_model or env.openai_model
    max_images = max(0, env.docx_vision_max_images)
    selected_images = image_names[:max_images] if max_images else []
    skipped = max(0, len(image_names) - len(selected_images))
    if not selected_images:
        return "", f"检测到 {len(image_names)} 张图片，但 DOCX_VISION_MAX_IMAGES=0，已跳过视觉模型解析。"

    try:
        from openai import OpenAI
    except ImportError:
        return extract_docx_images_with_ocr(path, image_names)

    client = OpenAI(api_key=api_key, base_url=base_url)
    image_blocks = []
    failed = 0
    with zipfile.ZipFile(path) as archive:
        for image_index, image_name in enumerate(selected_images, start=1):
            try:
                image_bytes = archive.read(image_name)
                data_url = image_to_data_url(image_name, image_bytes)
                completion = client.chat.completions.create(
                    model=model,
                    messages=[
                        {
                            "role": "system",
                            "content": "你是运维文档解析助手。请读取图片中的文字、命令、错误信息、配置、表格和架构关系，不要编造。",
                        },
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": (
                                        "请解析这张运维文档图片，输出 Markdown。"
                                        "如果是截图，提取界面文字、报错、命令、状态；"
                                        "如果是架构图，描述组件关系；如果看不清，请说明。"
                                    ),
                                },
                                {"type": "image_url", "image_url": {"url": data_url}},
                            ],
                        },
                    ],
                )
                text = (completion.choices[0].message.content or "").strip()
                if text:
                    image_blocks.append(f"### 图片 {image_index}: {Path(image_name).name}\n{text}")
            except Exception:
                failed += 1

    notes = []
    if skipped:
        notes.append(f"共有 {len(image_names)} 张图片，已按上限解析前 {len(selected_images)} 张，跳过 {skipped} 张。")
    if failed:
        notes.append(f"{failed} 张图片视觉模型解析失败。")
    if not image_blocks:
        notes.append("视觉模型未解析出有效图片内容。")
    return "\n\n".join(image_blocks), "\n".join(notes)


def image_to_data_url(image_name: str, image_bytes: bytes) -> str:
    suffix = Path(image_name).suffix.lower()
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".bmp": "image/bmp",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    mime = mime_map.get(suffix, "image/png")
    return f"data:{mime};base64,{base64.b64encode(image_bytes).decode('ascii')}"


def extract_docx_images_with_ocr(path: Path, image_names: list[str]) -> tuple[str, str]:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return "", f"检测到 {len(image_names)} 张图片，但当前环境未安装 pytesseract/Pillow，图片中的文字未识别。"

    if not tesseract_available():
        return "", f"检测到 {len(image_names)} 张图片，但当前环境没有 tesseract 命令，图片中的文字未识别。配置 OPENAI_API_KEY 后会优先使用视觉大模型解析图片。"

    image_blocks = []
    failed = 0
    with zipfile.ZipFile(path) as archive:
        for image_index, image_name in enumerate(image_names, start=1):
            try:
                image = Image.open(io.BytesIO(archive.read(image_name)))
                text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
                if text:
                    image_blocks.append(f"### 图片 {image_index}: {Path(image_name).name}\n{text}")
            except Exception:
                failed += 1

    note = ""
    if failed:
        note = f"{failed} 张图片 OCR 失败，可能是格式不支持、图片过小或缺少中文语言包。"
    if not image_blocks and not note:
        note = f"检测到 {len(image_names)} 张图片，但 OCR 未识别出文字。"
    return "\n\n".join(image_blocks), note


def tesseract_available() -> bool:
    import shutil

    return shutil.which("tesseract") is not None


def list_docx_images(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        return [
            name
            for name in archive.namelist()
            if name.startswith("word/media/")
            and Path(name).suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".webp"}
        ]


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


HEADING_RE = re.compile(r"^#{1,6}\s+")


def split_text(text: str, chunk_size: int = 900, overlap: int = 120) -> list[str]:
    if not text:
        return []
    text = clean_text(text)
    chunk_size = max(chunk_size, 1)
    overlap = min(max(overlap, 0), max(chunk_size - 1, 0))
    if len(text) <= chunk_size:
        return [text]

    blocks = _markdown_blocks(text)
    if len(blocks) <= 1:
        return _split_fixed_window(text, chunk_size, overlap)

    chunks: list[str] = []
    current = ""
    current_heading = ""

    for block in blocks:
        if _is_heading_block(block):
            current_heading = block.splitlines()[0].strip()

        pieces = _split_fixed_window(block, chunk_size, overlap) if len(block) > chunk_size else [block]
        for piece in pieces:
            piece = _with_heading_context(piece, current_heading, chunk_size)
            separator = "\n\n" if current else ""
            if current and len(current) + len(separator) + len(piece) > chunk_size:
                chunks.append(current.strip())
                current = _start_chunk_with_overlap(chunks[-1], piece, chunk_size, overlap)
            else:
                current = f"{current}{separator}{piece}" if current else piece

    if current.strip():
        chunks.append(current.strip())
    return [chunk for chunk in chunks if chunk]


def _markdown_blocks(text: str) -> list[str]:
    blocks: list[str] = []
    current: list[str] = []
    in_fence = False

    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_fence = not in_fence
            current.append(line.rstrip())
            continue
        if not in_fence and not stripped:
            if current:
                blocks.append("\n".join(current).strip())
                current = []
            continue
        if not in_fence and HEADING_RE.match(stripped):
            if current:
                blocks.append("\n".join(current).strip())
            current = [stripped]
            continue
        current.append(line.rstrip())

    if current:
        blocks.append("\n".join(current).strip())
    return [block for block in blocks if block]


def _split_fixed_window(text: str, chunk_size: int, overlap: int) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end].strip())
        if end == len(text):
            break
        start = max(start + 1, end - overlap)
    return [chunk for chunk in chunks if chunk]


def _is_heading_block(block: str) -> bool:
    return bool(HEADING_RE.match(block.splitlines()[0].strip())) if block else False


def _with_heading_context(piece: str, heading: str, chunk_size: int) -> str:
    if not heading or _is_heading_block(piece):
        return piece
    if piece.startswith(f"{heading}\n") or piece == heading:
        return piece
    candidate = f"{heading}\n\n{piece}"
    return candidate if len(candidate) <= chunk_size else piece


def _start_chunk_with_overlap(previous: str, piece: str, chunk_size: int, overlap: int) -> str:
    if not overlap:
        return piece
    available = chunk_size - len(piece) - 2
    if available < 40:
        return piece
    prefix = previous[-min(overlap, available):].strip()
    return f"{prefix}\n\n{piece}" if prefix else piece


def create_document_task(db: Session, document: Document, task_type: str = "parse") -> DocumentTask:
    task = DocumentTask(document_id=document.id, task_type=task_type, status="queued", progress=0, message="等待处理")
    db.add(task)
    db.commit()
    db.refresh(task)
    return task


def update_document_task(
    db: Session,
    task_id: int | None,
    status: str,
    progress: int | None = None,
    message: str = "",
) -> None:
    if not task_id:
        return
    task = db.get(DocumentTask, task_id)
    if not task:
        return
    task.status = status
    if progress is not None:
        task.progress = min(max(progress, 0), 100)
    if message:
        task.message = message
    db.commit()


def create_document_revision(db: Session, document: Document, markdown: str, note: str = "") -> DocumentRevision:
    latest = (
        db.query(func.max(DocumentRevision.version))
        .filter(DocumentRevision.document_id == document.id)
        .scalar()
        or 0
    )
    revision = DocumentRevision(document_id=document.id, version=latest + 1, content=markdown, note=note)
    db.add(revision)
    db.flush()
    return revision


def list_document_revisions(db: Session, document_id: int) -> list[DocumentRevision]:
    return (
        db.query(DocumentRevision)
        .filter(DocumentRevision.document_id == document_id)
        .order_by(DocumentRevision.version.desc())
        .all()
    )


def restore_document_revision(db: Session, document: Document, revision_id: int) -> None:
    revision = db.get(DocumentRevision, revision_id)
    if not revision or revision.document_id != document.id:
        raise ValueError("修订版本不存在")
    save_document_markdown(db, document, revision.content, note=f"restore-v{revision.version}")


def process_document(db: Session, document_id: int, task_id: int | None = None) -> None:
    document = db.get(Document, document_id)
    if not document:
        return
    try:
        update_document_task(db, task_id, "running", 10, "开始解析文档")
        document.status = "parsing"
        db.commit()
        workspace = document_workspace(document_id)
        shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        update_document_task(db, task_id, "running", 35, "提取正文和图片")
        markdown = build_markdown_document(Path(document.file_path), document_id, db)
        document_markdown_path(document_id).write_text(markdown, encoding="utf-8")
        update_document_task(db, task_id, "running", 70, "重建知识库索引")
        rebuild_document_chunks(db, document, markdown)
        create_document_revision(db, document, markdown, note="parsed")
        document.status = "completed"
        document.error_message = ""
        update_document_task(db, task_id, "completed", 100, "解析完成")
    except Exception as exc:
        document.status = "failed"
        document.error_message = str(exc)
        update_document_task(db, task_id, "failed", 100, str(exc))
    finally:
        db.commit()


def save_document_markdown(db: Session, document: Document, markdown: str, note: str = "manual-edit") -> None:
    workspace = document_workspace(document.id)
    workspace.mkdir(parents=True, exist_ok=True)
    document_markdown_path(document.id).write_text(markdown, encoding="utf-8")
    rebuild_document_chunks(db, document, markdown)
    create_document_revision(db, document, markdown, note=note)
    document.status = "completed"
    document.error_message = ""
    db.commit()


def create_knowledge_draft(
    db: Session,
    title: str,
    markdown: str,
    project: str = "default",
    note: str = "agent-draft",
) -> Document:
    """创建可检索的知识草稿，不覆盖正式文档。调用方负责 commit。"""
    document = Document(title=title.strip()[:255] or "Agent 知识草稿", file_path="", project=project, status="draft")
    db.add(document)
    db.flush()
    workspace = document_workspace(document.id)
    workspace.mkdir(parents=True, exist_ok=True)
    document.file_path = str(document_markdown_path(document.id))
    document_markdown_path(document.id).write_text(markdown, encoding="utf-8")
    rebuild_document_chunks(db, document, markdown)
    create_document_revision(db, document, markdown, note=note)
    document.error_message = ""
    db.flush()
    return document


def rebuild_document_chunks(db: Session, document: Document, markdown: str) -> None:
    chunks = split_text(clean_text(markdown))
    _ensure_fts(db)
    delete_document_fts(db, document.id)
    document.chunks.clear()
    db.flush()
    created_chunks: list[DocumentChunk] = []
    for index, chunk_text in enumerate(chunks):
        chunk = DocumentChunk(
            chunk_index=index,
            content=chunk_text,
            source=f"{document.title}#chunk-{index + 1}",
            project=document.project,
        )
        document.chunks.append(chunk)
        db.flush()
        _insert_fts(db, chunk)
        created_chunks.append(chunk)
    _upsert_chunk_embeddings(db, created_chunks)


def _raw_sql(db: Session, sql: str, params: tuple | dict | None = None):
    """执行原始 SQL，复用 Session 的连接（避免锁冲突）."""
    conn = db.connection()
    result = conn.exec_driver_sql(sql, params) if params else conn.exec_driver_sql(sql)
    return result


def _ensure_fts(db: Session) -> None:
    """确保 FTS5 全文索引表存在."""
    _raw_sql(db, "CREATE VIRTUAL TABLE IF NOT EXISTS chunk_fts USING fts5(content, project, tokenize='unicode61')")


def _insert_fts(db: Session, chunk: DocumentChunk) -> None:
    _raw_sql(
        db,
        "INSERT OR REPLACE INTO chunk_fts(rowid, content, project) VALUES (?, ?, ?)",
        (chunk.id, chunk.content, chunk.project),
    )


def delete_document_fts(db: Session, document_id: int) -> None:
    _ensure_fts(db)
    _raw_sql(
        db,
        "DELETE FROM chunk_fts WHERE rowid IN (SELECT id FROM document_chunks WHERE document_id = ?)",
        (document_id,),
    )


def search_chunks(db: Session, project: str, query: str, limit: int = 5) -> list[DocumentChunk]:
    """向量语义召回 + FTS5/LIKE 关键词召回的混合检索."""
    limit = min(max(limit, 1), 20)
    keyword_chunks = _keyword_search_chunks(db, project, query, limit)
    vector_hits = _vector_search_chunks(db, project, query, limit)
    if not vector_hits:
        return keyword_chunks
    return _merge_hybrid_results(keyword_chunks, vector_hits, limit)


def _keyword_search_chunks(db: Session, project: str, query: str, limit: int = 5) -> list[DocumentChunk]:
    """FTS5 + 标题/来源补充召回."""
    _ensure_fts(db)
    limit = min(max(limit, 1), 20)
    terms = re.findall(r"[\w\u4e00-\u9fff]+", query)
    if not terms:
        return db.query(DocumentChunk).filter(DocumentChunk.project == project).order_by(DocumentChunk.id.desc()).limit(limit).all()

    fts_query = _build_fts_query(terms)
    rows = []
    if fts_query:
        sql = (
            "SELECT rowid FROM chunk_fts "
            "WHERE chunk_fts MATCH :q AND project = :proj "
            "ORDER BY rank LIMIT :lim"
        )
        try:
            result = _raw_sql(db, sql, {"q": fts_query, "proj": project, "lim": limit})
            rows = result.fetchall()
        except Exception:
            rows = []
    ids = [r[0] for r in rows]
    ranked: list[DocumentChunk] = []
    seen: set[int] = set()
    if ids:
        order = {chunk_id: idx for idx, chunk_id in enumerate(ids)}
        ranked = db.query(DocumentChunk).filter(DocumentChunk.id.in_(ids)).all()
        ranked.sort(key=lambda chunk: order.get(chunk.id, 9999))
        seen.update(chunk.id for chunk in ranked)
    filters = []
    for term in terms[:5]:
        if len(term) > 1:
            like = f"%{term}%"
            filters.append(DocumentChunk.content.like(like))
            filters.append(DocumentChunk.source.like(like))
            filters.append(Document.title.like(like))
    if not filters:
        if not ranked:
            return db.query(DocumentChunk).filter(DocumentChunk.project == project).order_by(DocumentChunk.id.desc()).limit(limit).all()
        return ranked[:limit]
    fallback = (
        db.query(DocumentChunk)
        .join(Document, Document.id == DocumentChunk.document_id)
        .filter(DocumentChunk.project == project)
        .filter(or_(*filters))
        .order_by(DocumentChunk.id.desc())
        .limit(limit * 4)
        .all()
    )
    fallback.sort(key=lambda chunk: (_fallback_score(chunk, terms), chunk.id), reverse=True)
    for chunk in fallback:
        if chunk.id not in seen:
            ranked.append(chunk)
            seen.add(chunk.id)
        if len(ranked) >= limit:
            break
    return ranked[:limit]


def _sanitize_fts(query: str) -> str:
    """清洗 FTS5 查询，移除特殊字符."""
    return re.sub(r'[^\w\u4e00-\u9fff\s]', ' ', query).strip()


def _build_fts_query(terms: list[str]) -> str:
    reserved = {"AND", "OR", "NOT", "NEAR"}
    phrases = []
    seen = set()
    for term in terms:
        cleaned = _sanitize_fts(term)
        if len(cleaned) <= 1 or cleaned.upper() in reserved or cleaned in seen:
            continue
        phrases.append(f'"{cleaned}"')
        seen.add(cleaned)
    return " OR ".join(phrases)


def _fallback_score(chunk: DocumentChunk, terms: list[str]) -> int:
    document = chunk.document
    title = (document.title if document else "").lower()
    source = (chunk.source or "").lower()
    content = (chunk.content or "").lower()
    score = 0
    for term in terms[:5]:
        needle = term.lower()
        if not needle:
            continue
        if title == needle:
            score += 20
        if needle in title:
            score += 12
        if needle in source:
            score += 6
        position = content.find(needle)
        if position >= 0:
            score += 4
            score += max(0, 3 - position // 300)
    return score


def _resolve_embedding_config(db: Session) -> dict:
    env = get_settings()
    keys = [
        "embedding_api_key",
        "embedding_base_url",
        "embedding_model",
        "chat_api_key",
        "chat_base_url",
    ]
    db_vals = {}
    for row in db.query(AppSetting).filter(AppSetting.key.in_(keys)).all():
        db_vals[row.key] = row.value
    return {
        "api_key": db_vals.get("embedding_api_key") or db_vals.get("chat_api_key") or getattr(env, "openai_api_key", "") or "",
        "base_url": db_vals.get("embedding_base_url") or db_vals.get("chat_base_url") or getattr(env, "openai_base_url", None) or None,
        "model": db_vals.get("embedding_model") or getattr(env, "openai_embedding_model", "text-embedding-3-small") or "text-embedding-3-small",
    }


def _upsert_chunk_embeddings(db: Session, chunks: list[DocumentChunk]) -> None:
    if not chunks:
        return
    config = _resolve_embedding_config(db)
    if not config["api_key"]:
        return

    texts = [_embedding_text(chunk) for chunk in chunks]
    try:
        vectors = _embed_texts(texts, config)
    except Exception:
        return
    if len(vectors) != len(chunks):
        return

    model = config["model"]
    for chunk, text, vector in zip(chunks, texts, vectors, strict=False):
        if not vector:
            continue
        payload = json.dumps(vector, separators=(",", ":"))
        content_hash = _embedding_hash(text)
        if chunk.embedding:
            embedding = chunk.embedding
            embedding.model = model
            embedding.dimension = len(vector)
            embedding.content_hash = content_hash
            embedding.vector = payload
        else:
            chunk.embedding = DocumentChunkEmbedding(
                chunk_id=chunk.id,
                model=model,
                dimension=len(vector),
                content_hash=content_hash,
                vector=payload,
            )
    db.flush()


def _vector_search_chunks(db: Session, project: str, query: str, limit: int) -> list[tuple[DocumentChunk, float]]:
    config = _resolve_embedding_config(db)
    if not config["api_key"]:
        return []
    try:
        query_vector = _embed_texts([query], config)[0]
    except Exception:
        return []
    if not query_vector:
        return []

    rows = (
        db.query(DocumentChunkEmbedding, DocumentChunk)
        .join(DocumentChunk, DocumentChunk.id == DocumentChunkEmbedding.chunk_id)
        .filter(DocumentChunk.project == project)
        .filter(DocumentChunkEmbedding.model == config["model"])
        .all()
    )
    hits: list[tuple[DocumentChunk, float]] = []
    for embedding, chunk in rows:
        vector = _decode_vector(embedding.vector)
        if not vector or embedding.dimension != len(query_vector):
            continue
        score = _cosine_similarity(query_vector, vector)
        if score > 0:
            hits.append((chunk, score))
    hits.sort(key=lambda item: (item[1], item[0].id), reverse=True)
    return hits[: max(limit * 4, limit)]


def _merge_hybrid_results(
    keyword_chunks: list[DocumentChunk],
    vector_hits: list[tuple[DocumentChunk, float]],
    limit: int,
) -> list[DocumentChunk]:
    scores: dict[int, float] = {}
    chunks: dict[int, DocumentChunk] = {}

    for rank, chunk in enumerate(keyword_chunks):
        chunks[chunk.id] = chunk
        scores[chunk.id] = scores.get(chunk.id, 0.0) + 0.35 / (rank + 1)

    for rank, (chunk, vector_score) in enumerate(vector_hits):
        chunks[chunk.id] = chunk
        semantic = max(vector_score, 0.0)
        scores[chunk.id] = scores.get(chunk.id, 0.0) + semantic * 0.75 + 0.15 / (rank + 1)

    ordered_ids = sorted(scores, key=lambda chunk_id: (scores[chunk_id], chunks[chunk_id].id), reverse=True)
    return [chunks[chunk_id] for chunk_id in ordered_ids[:limit]]


def _embed_texts(texts: list[str], config: dict) -> list[list[float]]:
    from openai import OpenAI

    client = OpenAI(api_key=config["api_key"], base_url=config["base_url"], timeout=45.0)
    response = client.embeddings.create(model=config["model"], input=texts)
    ordered = sorted(response.data, key=lambda item: getattr(item, "index", 0))
    return [list(item.embedding) for item in ordered]


def _embedding_text(chunk: DocumentChunk) -> str:
    return f"{chunk.source}\n\n{chunk.content}".strip()


def _embedding_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _decode_vector(payload: str) -> list[float]:
    try:
        data = json.loads(payload)
    except Exception:
        return []
    if not isinstance(data, list):
        return []
    try:
        return [float(value) for value in data]
    except (TypeError, ValueError):
        return []


def _cosine_similarity(left: list[float], right: list[float]) -> float:
    if len(left) != len(right) or not left:
        return 0.0
    dot = sum(a * b for a, b in zip(left, right, strict=False))
    left_norm = math.sqrt(sum(a * a for a in left))
    right_norm = math.sqrt(sum(b * b for b in right))
    if not left_norm or not right_norm:
        return 0.0
    return dot / (left_norm * right_norm)
